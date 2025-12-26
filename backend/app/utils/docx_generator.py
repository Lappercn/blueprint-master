# 文件名：docx_generator.py
"""
功能说明：Markdown 转 Word 文档生成器
核心功能：解析 Markdown 文本并生成格式化的 .docx 文件
依赖模块：python-docx
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import re
import io

def generate_blueprint_docx(markdown_text: str) -> io.BytesIO:
    """
    将 Markdown 文本转换为 Word 文档流
    :param markdown_text: Markdown 格式的分析报告
    :return: BytesIO 对象 (docx 文件流)
    """
    document = Document()
    
    # 设置中文字体
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 1. 标题处理
        if line.startswith('# '):
            # 一级标题
            h1 = document.add_heading(line.replace('# ', '').strip(), level=1)
            run = h1.runs[0]
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            run.font.color.rgb = RGBColor(216, 30, 6) # 红色
        elif line.startswith('## '):
            # 二级标题
            h2 = document.add_heading(line.replace('## ', '').strip(), level=2)
            run = h2.runs[0]
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('### '):
            # 三级标题
            h3 = document.add_heading(line.replace('### ', '').strip(), level=3)
            run = h3.runs[0]
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            
        # 2. 引用块处理
        elif line.startswith('> '):
            p = document.add_paragraph()
            run = p.add_run(line.replace('> ', '').strip())
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
        # 3. 列表项处理
        elif line.startswith('* ') or line.startswith('- '):
            content = line[2:].strip()
            p = document.add_paragraph(style='List Bullet')
            _add_formatted_run(p, content)
            
        # 4. 数字列表
        elif re.match(r'^\d+\.\s', line):
            content = re.sub(r'^\d+\.\s', '', line).strip()
            p = document.add_paragraph(style='List Number')
            _add_formatted_run(p, content)
            
        # 5. 普通段落
        else:
            p = document.add_paragraph()
            _add_formatted_run(p, line)
            
    # 保存到内存流
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def _add_formatted_run(paragraph, text):
    """
    处理段落中的粗体 (**text**)
    """
    # 简单的正则匹配粗体
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        else:
            paragraph.add_run(part)
